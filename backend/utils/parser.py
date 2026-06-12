from io import BytesIO

from fastapi import UploadFile


async def extract_text(file: UploadFile) -> str:
    content = await file.read()
    filename = (file.filename or "").lower()

    if filename.endswith(".pdf"):
        return _extract_pdf(content)
    if filename.endswith(".docx") or filename.endswith(".doc"):
        return _extract_docx(content)
    if any(filename.endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tiff")):
        return _extract_image(content)

    try:
        return content.decode("utf-8").strip()
    except UnicodeDecodeError:
        return content.decode("latin-1", errors="replace").strip()


def _extract_pdf(content: bytes) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(stream=content, filetype="pdf")
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(pages).strip()
    except Exception as e:
        return f"[PDF extraction failed: {e}]"


def _extract_docx(content: bytes) -> str:
    try:
        import docx
        doc = docx.Document(BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)
        return "\n".join(paragraphs).strip()
    except Exception as e:
        return f"[DOCX extraction failed: {e}]"


def _extract_image(content: bytes) -> str:
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(BytesIO(content))
        return pytesseract.image_to_string(img).strip()
    except Exception as e:
        return f"[Image OCR failed: {e}]"
